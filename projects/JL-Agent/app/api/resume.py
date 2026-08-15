"""Resume CRUD（契约 §4.2）：POST/GET/PUT/DELETE + 列表 + 条目编辑锁定（§5.5）+ 重装配渲染（§6）+ 导出（§7 E8）。"""
import json
from typing import Optional

from fastapi import APIRouter, Depends, Request
from fastapi.responses import Response
from pydantic import BaseModel, Field

from ..core.errors import AppError, E_EXPORT, E_PARAM
from ..core.validation import check_resume
from ..engine.assembly import Assembler
from ..schemas import CamelModel, Resume

router = APIRouter(prefix="/api/resume", tags=["resume"])

# 可编辑条目块（§5.5）：summary 句 / internship duty / project item
EDITABLE_BLOCKS = ("summary", "internship", "project")


class ResumeIdResp(BaseModel):
    resume_id: str


class DeletedResp(BaseModel):
    deleted: bool


class ResumeListItem(BaseModel):
    id: str
    name: str = ""
    direction: str = ""
    updated_at: Optional[str] = None
    file: str = ""     # 本地存储位置（相对路径，供 UI 展示）


class ItemEditBody(CamelModel):
    """编辑锁定（§5.5）：summary 用 index；internship/project 用 index + sub_index 定位叶子。"""
    block: str = Field(min_length=1)                       # summary/internship/project
    index: int = Field(ge=0)                               # 父级下标（summary 为句子下标）
    sub_index: Optional[int] = Field(default=None, ge=0)   # 叶子下标（实习职责/项目要点）
    text: str = Field(min_length=1, max_length=500)


class ItemUnlockBody(CamelModel):
    block: str = Field(min_length=1)
    index: int = Field(ge=0)
    sub_index: Optional[int] = Field(default=None, ge=0)


class RenderBody(BaseModel):
    density: Optional[str] = Field(default=None, pattern="^(compact|normal|loose)$")


@router.post("", response_model=dict)
def create_resume(body: Resume, request: Request):
    """新建简历：无 id 则生成；集中校验后落库。"""
    storage = request.app.state.storage
    now = request.app.state.now
    if not body.id:
        body.id = storage.new_resume_id()
    body.created_at = body.created_at or now()
    body.updated_at = now()
    check_resume(body, request.app.state.config.limits)
    storage.save_resume(body.model_dump(mode="json", by_alias=True, exclude_none=False))
    return {"code": 0, "message": "ok", "data": {"resumeId": body.id}}


@router.get("", response_model=dict)
def list_resumes(request: Request):
    """简历列表（轻量）：前端工作台使用。"""
    storage = request.app.state.storage
    items = []
    for rid in storage.list_resumes():
        data = storage.load_resume(rid)
        name = (data.get("basicInfo") or {}).get("name", "")
        direction = data.get("direction") or ""
        items.append(ResumeListItem(id=rid, name=name, direction=direction,
                                    updated_at=data.get("updatedAt"),
                                    file=f"data/resumes/{rid}.json").model_dump())
    items.sort(key=lambda x: x["updated_at"] or "", reverse=True)
    return {"code": 0, "message": "ok", "data": {"items": items}}


@router.get("/{resume_id}", response_model=dict)
def get_resume(resume_id: str, request: Request):
    storage = request.app.state.storage
    data = storage.load_resume(resume_id)
    return {"code": 0, "message": "ok", "data": data}


@router.put("/{resume_id}", response_model=dict)
def update_resume(resume_id: str, body: Resume, request: Request):
    """整存更新：保留 id/created_at，刷新 updated_at。"""
    storage = request.app.state.storage
    now = request.app.state.now
    old = storage.load_resume(resume_id)
    body.id = resume_id
    body.created_at = body.created_at or old.get("created_at")
    body.updated_at = now()
    check_resume(body, request.app.state.config.limits)
    storage.save_resume(body.model_dump(mode="json", by_alias=True, exclude_none=False))
    return {"code": 0, "message": "ok", "data": {"updatedAt": body.updated_at}}


@router.delete("/{resume_id}", response_model=dict)
def delete_resume(resume_id: str, request: Request):
    storage = request.app.state.storage
    storage.load_resume(resume_id)  # 不存在 → 40008
    deleted = storage.delete_resume(resume_id)
    return {"code": 0, "message": "ok", "data": DeletedResp(deleted=deleted).model_dump()}


# ---------------------------------------------------------------- 编辑锁定（§5.5）与重装配（§6）

def _leaf(resume: dict, block: str, index: int, sub_index: Optional[int]) -> dict:
    """定位可编辑叶子条目：summary 句 / 实习职责 / 项目要点。非法定位 → 40001。"""
    if block == "summary":
        if sub_index is not None:
            raise AppError(E_PARAM, "summary 是单层列表，无需 subIndex", {"block": block})
        items = resume.get("summary") or []
        idx = index
    elif block in ("internship", "project"):
        parents = resume.get(block) or []
        if not (0 <= index < len(parents)):
            raise AppError(E_PARAM, f"{block} 下标越界", {"block": block, "index": index})
        key = "duties" if block == "internship" else "items"
        items = parents[index].get(key) or []
        if sub_index is None:
            raise AppError(E_PARAM, f"{block} 需要 subIndex 定位叶子", {"block": block})
        idx = sub_index
    else:
        raise AppError(E_PARAM, "不可编辑板块", {"block": block})
    if not (0 <= idx < len(items)):
        raise AppError(E_PARAM, "条目下标越界",
                       {"block": block, "index": index, "subIndex": sub_index})
    return items[idx]


def _rendered(resume: dict, app) -> dict:
    """重装配（§6）：density 取 resume 现值，返回 {resume, html, config} 供预览。"""
    assembler = Assembler(app.state.config.paths.templates_dir, app.state.storage)
    generation = resume.get("generation") or {}
    html, config = assembler.render(
        resume, {},
        density=str(resume.get("density") or "normal"),
        watermark_mode=str(generation.get("watermarkMode") or "practice"),
    )
    return {"code": 0, "message": "ok",
            "data": {"resume": resume, "html": html, "config": config}}


@router.put("/{resume_id}/item", response_model=dict)
def edit_item(resume_id: str, body: ItemEditBody, request: Request):
    """单条编辑（§5.5）：改文本 + edited:true + criticality 强制 critical → 不可被自动重写。"""
    storage = request.app.state.storage
    resume = storage.load_resume(resume_id)
    leaf = _leaf(resume, body.block, body.index, body.sub_index)
    text = body.text.strip()
    if not text:
        raise AppError(E_PARAM, "文本不能为空", {"block": body.block})
    leaf["text"] = text[: (300 if body.block == "summary" else 500)]
    leaf["edited"] = True
    leaf["criticality"] = "critical"
    resume["updatedAt"] = request.app.state.now()
    storage.save_resume(resume)
    return _rendered(resume, request.app)


@router.post("/{resume_id}/item/unlock", response_model=dict)
def unlock_item(resume_id: str, body: ItemUnlockBody, request: Request):
    """解锁（§5.5）：edited:false → 下次自动生成可重写该条目。"""
    storage = request.app.state.storage
    resume = storage.load_resume(resume_id)
    leaf = _leaf(resume, body.block, body.index, body.sub_index)
    leaf["edited"] = False
    resume["updatedAt"] = request.app.state.now()
    storage.save_resume(resume)
    return _rendered(resume, request.app)


@router.post("/{resume_id}/render", response_model=dict)
def render_resume(resume_id: str, body: RenderBody, request: Request):
    """重装配渲染（§6）：支持 density 手动调整后重新出图。"""
    storage = request.app.state.storage
    resume = storage.load_resume(resume_id)
    if body.density:
        resume["density"] = body.density
        resume["updatedAt"] = request.app.state.now()
        storage.save_resume(resume)
    return _rendered(resume, request.app)


# ---------------------------------------------------------------- 导出（§7 E8：PDF / DOCX / JSON）


def _contact_line(info: dict) -> str:
    """个人信息一行（带标签，如「邮箱：qixiao.e@testmail.com」）。"""
    segs = []
    if info.get("phone"):
        segs.append(f"电话：{info['phone']}")
    if info.get("email"):
        segs.append(f"邮箱：{info['email']}")
    if info.get("website"):
        segs.append(f"个人网站：{info['website']}")
    if info.get("base"):
        segs.append(f"所在城市：{info['base']}")
    if info.get("internshipDuration"):
        segs.append(f"可实习时长：{info['internshipDuration']}")
    if info.get("startAvailable"):
        segs.append(f"到岗时间：{info['startAvailable']}")
    return "  ".join(segs)


def _structured_sections(resume: dict) -> list:
    """结构化区块（docx/md 复用）：目标顺序 教育→荣誉→实习→项目→技能→个人评价，
    含时间段 / 技术栈 / 技能分类行。"""
    blocks = []

    def add(title: str, lines: list) -> None:
        lines = [str(x).strip() for x in lines if x is not None and str(x).strip()]
        if lines:
            blocks.append((title, lines))

    add("教育经历", [
        f"{e.get('school')} · {e.get('major')}（{e.get('degree')}）"
        f"{'  ' + e.get('startMonth') + '—' + e.get('endMonth') if e.get('startMonth') or e.get('endMonth') else ''}"
        for e in (resume.get("education") or [])])
    add("证书荣誉", [
        f"{h.get('name')}" + (f" · {h.get('time')}" if h.get("time") else "")
        for h in (resume.get("honor") or [])])
    int_lines = []
    for it in (resume.get("internship") or []):
        head = f"{it.get('company')} · {it.get('position')}"
        if it.get("startMonth") or it.get("endMonth"):
            head += f"  {it.get('startMonth')}—{it.get('endMonth')}"
        int_lines.append(head)
        if str(it.get("overview") or "").strip():
            int_lines.append("主要职责：" + str(it["overview"]).strip())
        int_lines += [d.get("text") for d in (it.get("duties") or [])]
    add("实习经历", int_lines)
    proj_lines = []
    for p in (resume.get("project") or []):
        head = str(p.get("name") or "")
        if p.get("startMonth") or p.get("endMonth"):
            head += f"  {p.get('startMonth')}—{p.get('endMonth')}"
        proj_lines.append(head)
        tech = "、".join(str(t) for t in (p.get("techStack") or []) if str(t).strip())
        if tech:
            proj_lines.append("技术栈：" + tech)
        proj_lines += [x.get("text") for x in (p.get("items") or [])]
    add("项目经验", proj_lines)
    groups = {}
    for s in (resume.get("skill") or []):
        name = str(s.get("name") or "").strip()
        if not name:
            continue
        groups.setdefault(str(s.get("category") or "其他"), []).append(name)
    add("技能特长", [f"{c}：{'、'.join(names)}" for c, names in groups.items()])
    add("个人评价", [s.get("text") for s in (resume.get("summary") or [])])
    return blocks


def _resume_to_docx(resume: dict) -> bytes:
    """结构化 DOCX：个人信息（带标签）+ 目标板块顺序 + 时间段 / 技术栈 / 技能分类（python-docx）。"""
    from io import BytesIO

    import docx

    document = docx.Document()
    info = resume.get("basicInfo") or {}
    document.add_heading(info.get("name") or "简历", level=0)
    contact = _contact_line(info)
    if contact:
        document.add_paragraph(contact)

    for title, lines in _structured_sections(resume):
        document.add_heading(title, level=1)
        for ln in lines:
            document.add_paragraph(ln, style="List Bullet")

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def _resume_to_markdown(resume: dict) -> str:
    """导出 Markdown：个人信息（带标签）+ 目标板块顺序（标准 GFM）。"""
    out = []
    info = resume.get("basicInfo") or {}
    out.append("# " + (info.get("name") or "简历"))
    contact = _contact_line(info)
    if contact:
        out.append("")
        out.append(contact)
    out.append("")

    for title, lines in _structured_sections(resume):
        out.append("## " + title)
        out.extend("- " + ln for ln in lines)
        out.append("")
    return "\n".join(out).strip() + "\n"


def _resume_to_html(resume: dict, templates_dir: str, storage) -> str:
    """导出完整 HTML5：复用模板装配（与预览一致的 A4 打印版式，含 @page 与打印样式）。"""
    assembler = Assembler(templates_dir, storage)
    generation = resume.get("generation") or {}
    html, _ = assembler.render(
        resume, {},
        density=str(resume.get("density") or "normal"),
        watermark_mode=str(generation.get("watermarkMode") or "formal"),
    )
    return html


@router.get("/{resume_id}/export", response_model=dict)
def export_resume(resume_id: str, request: Request, format: str = "json"):
    """导出：format=json（结构化数据）/ docx（Word）/ md|markdown / html / pdf（由前端打印生成）。"""
    storage = request.app.state.storage
    resume = storage.load_resume(resume_id)
    fmt = format.lower()
    if fmt == "json":
        data = json.dumps(resume, ensure_ascii=False, indent=2).encode("utf-8")
        return Response(
            content=data, media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{resume_id}.json"'})
    if fmt == "docx":
        try:
            content = _resume_to_docx(resume)
        except ImportError:
            raise AppError(E_EXPORT, "DOCX 导出需要安装 python-docx（pip install python-docx）")
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{resume_id}.docx"'})
    if fmt in ("md", "markdown"):
        content = _resume_to_markdown(resume).encode("utf-8")
        return Response(
            content=content, media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{resume_id}.md"'})
    if fmt == "html":
        content = _resume_to_html(resume, request.app.state.config.paths.templates_dir, storage).encode("utf-8")
        return Response(
            content=content, media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{resume_id}.html"'})
    raise AppError(E_PARAM, "不支持的导出格式", {"format": fmt})
